# import os
# import re
# import docx
# import pdfplumber
# import textract
# from pdf2image import convert_from_path
# from pytesseract import image_to_string
# from win32com.client import Dispatch

# # ==============================
# # Extract Text from Different File Formats
# # ==============================

# def extract_text_from_pdf(file_path):
#     """Extract text from a PDF file, with OCR fallback for scanned PDFs."""
#     text = ""
#     try:
#         with pdfplumber.open(file_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text() or ""
#                 text += page_text + "\n"
#         if not text.strip():
#             print("No text extracted, falling back to OCR.")
#             text = perform_ocr(file_path)
#     except Exception as e:
#         print(f"Error reading PDF file {file_path}: {e}")
#     return text.strip()

# def extract_text_from_docx(file_path):
#     """Extract text from a DOCX file, including paragraphs, tables, headers, and footers."""
#     text = []
#     try:
#         doc = docx.Document(file_path)

#         # Extract paragraphs
#         for para in doc.paragraphs:
#             text.append(para.text.strip())

#         # Extract tables
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 text.append(row_text)

#         # Extract headers and footers
#         for section in doc.sections:
#             if section.header:
#                 for paragraph in section.header.paragraphs:
#                     text.append(paragraph.text.strip())
#             if section.footer:
#                 for paragraph in section.footer.paragraphs:
#                     text.append(paragraph.text.strip())

#     except Exception as e:
#         print(f"Error reading DOCX file {file_path}: {e}")

#     return "\n".join(filter(None, text)).strip()

# def extract_text_from_doc(file_path):
#     """Extract text from a DOC file using PyWin32 and Textract fallback."""
#     text = ""
#     try:
#         word = Dispatch("Word.Application")
#         word.Visible = False
#         doc = word.Documents.Open(file_path)
#         text = doc.Content.Text
#         doc.Close()
#         word.Quit()
#     except Exception as e:
#         print(f"PyWin32 failed for DOC file {file_path}: {e}")
#         try:
#             text = textract.process(file_path, extension="doc").decode("utf-8").strip()
#         except Exception as textract_error:
#             print(f"Textract also failed: {textract_error}")
#     return text.strip()

# # Perform OCR on image-based documents
# def perform_ocr(file_path):
#     try:
#         images = convert_from_path(file_path)
#         text = "\n".join(image_to_string(image, lang="eng") for image in images)
#         return text.strip()
#     except Exception as e:
#         print(f"OCR failed: {e}")
#         return ""

# # ==============================
# # Extract Emails from Text
# # ==============================

# def extract_emails(text):
#     """Extract email addresses and clean up output."""
#     email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
#     emails = set(re.findall(email_pattern, text))  # Ensure unique emails
#     return [email.strip() for email in emails if email.strip()]  # Remove empty values

# # ==============================
# # Process Resume Files
# # ==============================

# def parse_resume(file_path):
#     """Parse the resume and extract full text and emails."""
#     full_text = ""
    
#     if file_path.endswith('.pdf'):
#         full_text = extract_text_from_pdf(file_path)
#     elif file_path.endswith('.docx'):
#         full_text = extract_text_from_docx(file_path)
#     elif file_path.endswith('.doc'):
#         full_text = extract_text_from_doc(file_path)
#     else:
#         raise ValueError("Unsupported file format.")

#     emails = extract_emails(full_text)

#     # Print extracted emails and full text separately
#     print("\n" + "=" * 80)
#     print(f"Processing File: {os.path.basename(file_path)}")
#     print("=" * 80)
#     print(f"Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}")
#     print("\nExtracted Full Text:\n" + "=" * 80)
#     print(full_text if full_text else "No text extracted.")
#     print("=" * 80 + "\n")

# # ==============================
# # Process Multiple Resumes in a Folder
# # ==============================

# def process_folder(folder_path):
#     """Process all resumes in a folder and print extracted data."""
#     for filename in os.listdir(folder_path):
#         file_path = os.path.join(folder_path, filename)
#         if filename.endswith(('.pdf', '.docx', '.doc')):
#             try:
#                 parse_resume(file_path)
#             except Exception as e:
#                 print(f"Error processing file {filename}: {e}")

# # ==============================
# # Main Execution
# # ==============================

# if __name__ == "__main__":
#     folder_path = r"C:\pdf_extract\1Chief_Officer_Computing-tgoodwyn.pdf"
#     process_folder(folder_path)
##################################
# import os
# import re
# import subprocess
# import pdfplumber
# import docx
# import textwrap
# from pdf2image import convert_from_path
# from pytesseract import image_to_string

# # ============================
# # Convert DOC to DOCX
# # ============================
# def convert_doc_to_docx(file_path):
#     """Converts a .doc file to .docx format using LibreOffice."""
#     output_dir = os.path.dirname(file_path)
#     new_file = os.path.splitext(file_path)[0] + ".docx"
#     soffice_cmd = "/usr/bin/soffice"

#     if not os.path.exists(soffice_cmd):
#         print("❌ LibreOffice (`soffice`) not found! Install it with `sudo apt install libreoffice`")
#         return None
#     try:
#         subprocess.run(
#             [soffice_cmd, "--headless", "--convert-to", "docx", file_path, "--outdir", output_dir],
#             stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#         )
#         return new_file if os.path.exists(new_file) else None
#     except subprocess.CalledProcessError as e:
#         print(f"❌ Error converting {file_path}: {e.stderr}")
#         return None

# # ============================
# # Extract Text from PDF
# # ============================
# def extract_text_from_pdf(pdf_path):
#     """Extracts text from a PDF while preserving formatting."""
#     text = []
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text.append(page_text)

#         # If no text is extracted, use OCR
#         if not text:
#             print(f"🔍 No direct text found in {pdf_path}, performing OCR...")
#             text.append(perform_ocr(pdf_path))

#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
    
#     return format_resume_text("\n".join(text))

# # ============================
# # Extract Text from DOCX with Formatting
# # ============================
# def extract_text_from_docx(file_path):
#     """Extracts text from a DOCX file while preserving layout and formatting."""
#     text = []
#     try:
#         doc = docx.Document(file_path)

#         # Extract headers
#         for section in doc.sections:
#             if section.header:
#                 text.extend(f"📌 {p.text.strip()}" for p in section.header.paragraphs if p.text.strip())

#         # Extract paragraphs and preserve formatting
#         for para in doc.paragraphs:
#             stripped_text = para.text.strip()
#             if not stripped_text:
#                 continue
#             if para.style and para.style.name.startswith("Heading"):
#                 text.append(f"\n📢 {stripped_text.upper()}\n")  # Preserve headings
#             elif para.text.startswith(("•", "-", "*")):
#                 text.append(f"  ➜ {stripped_text}")  # Maintain bullet points
#             else:
#                 text.append(stripped_text)

#         # Extract tables and format them as structured text
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 if row_text:
#                     text.append(row_text)

#         # Extract footers
#         for section in doc.sections:
#             if section.footer:
#                 text.extend(f"📌 {p.text.strip()}" for p in section.footer.paragraphs if p.text.strip())
#     except Exception as e:
#         print(f"❌ Error reading DOCX file {file_path}: {e}")
#     return format_resume_text("\n".join(text))

# # ============================
# # Perform OCR on Scanned PDFs
# # ============================
# def perform_ocr(file_path):
#     """Performs OCR on scanned PDFs and extracts text for hidden emails."""
#     try:
#         images = convert_from_path(file_path)
#         text = "\n".join(image_to_string(img, lang="eng") for img in images).strip()
#         return text
#     except Exception as e:
#         print(f"❌ OCR failed: {e}")
#         return ""

# # ============================
# # Format Resume Text Properly
# # ============================
# def format_resume_text(text, width=100):
#     """Formats extracted text to improve layout and readability."""
#     lines = text.split("\n")
#     formatted_lines = []
#     is_heading = False

#     for line in lines:
#         stripped_line = line.strip()

#         # Handle Bullet Points
#         if stripped_line.startswith(("•", "-", "*")):
#             formatted_lines.append(f"  ➜ {stripped_line}")

#         # Preserve Section Headers
#         elif stripped_line.isupper():
#             formatted_lines.append(f"\n📢 {stripped_line}\n")
#             is_heading = True

#         # Ensure Proper Paragraph Formatting
#         elif is_heading:
#             formatted_lines.append(stripped_line)
#             is_heading = False
#         else:
#             formatted_lines.append(textwrap.fill(stripped_line, width=width))

#     return "\n".join(formatted_lines)

# # ============================
# # Extract Emails
# # ============================
# def extract_emails(text):
#     """Extracts email addresses from the text, including 'mailto:' links and multi-line formats."""
#     text = text.replace("\n", " ").replace("\r", " ")  # Normalize new lines and spaces
#     email_pattern = r"mailto:([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})|([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})"
    
#     emails = [match[0] if match[0] else match[1] for match in re.findall(email_pattern, text)]
#     return list(set(emails))  # Remove duplicates

# # ============================
# # Extract URLs (LinkedIn, Websites)
# # ============================
# def extract_urls(text):
#     """Extracts URLs (LinkedIn, GitHub, personal websites) from text."""
#     url_pattern = r"(https?://[^\s]+)"
#     urls = re.findall(url_pattern, text)
#     return list(set(urls))

# # ============================
# # Extract Resume Text
# # ============================
# def extract_resume_text(file_path):
#     """Extracts text from PDF, DOCX, or converts DOC to DOCX before extracting."""
#     if file_path.lower().endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path.lower().endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     elif file_path.lower().endswith(".doc"):
#         converted = convert_doc_to_docx(file_path)
#         return extract_text_from_docx(converted) if converted else ""
#     else:
#         raise ValueError("❌ Unsupported file format. Only PDF, DOC, and DOCX are supported.")

# # ============================
# # Process Resumes in Folder
# # ============================
# def process_resumes_in_folder(folder_path, limit=10):
#     """Processes up to limit resume files in the folder and extracts emails."""
#     files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".doc", ".docx"))][:limit]
#     if not files:
#         print("❌ No resume files found in the folder.")
#         return
    
#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         print("=" * 100)
#         print(f"📂 Processing File: {filename}")
#         print("=" * 100)
#         try:
#             text = extract_resume_text(file_path)
#             emails = extract_emails(text)
#             urls = extract_urls(text)

#             print(f"📧 Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}")
#             print(f"🔗 Extracted URLs: {', '.join(urls) if urls else 'No URLs found.'}\n")
#             print(text[:2000])  # Limit output

#         except Exception as e:
#             print(f"❌ Error processing file {filename}: {e}")
#         print("=" * 100 + "\n")

# # ============================
# # Main Execution
# # ============================
# if __name__ == "__main__":
#     folder_path = "/app"
#     process_resumes_in_folder(folder_path, limit=10)
####################################
import os
import re
import docx
import pdfplumber
import textract
from pdf2image import convert_from_path
from pytesseract import image_to_string

# ==============================
# Extract Text from Different File Formats
# ==============================

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file, with OCR fallback for scanned PDFs."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        # Fallback to OCR if no text was extracted
        if not text.strip():
            print("No text extracted from PDF, using OCR...")
            text = perform_ocr(file_path)
    
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
    
    return normalize_text(text)

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
    
    return normalize_text(text)

def perform_ocr(file_path):
    """Perform OCR on a scanned PDF file."""
    try:
        images = convert_from_path(file_path)
        text = "\n".join(image_to_string(image, lang="eng") for image in images)
        return normalize_text(text)
    except Exception as e:
        print(f"OCR failed: {e}")
        return ""

# ==============================
# Extract Emails from Text
# ==============================

def extract_emails(text):
    """Extract email addresses from text after normalization."""
    text = normalize_text(text)  # Ensure text is formatted correctly
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    emails = set(re.findall(email_pattern, text))  # Unique emails
    return list(emails) if emails else ["No emails found."]

# ==============================
# Text Normalization
# ==============================

def normalize_text(text):
    """Cleans extracted text by removing excessive spaces, special characters, and formatting it properly."""
    text = text.replace("\n", " ")  # Convert newlines to spaces
    text = re.sub(r'\s+', ' ', text)  # Remove excessive spaces
    text = text.strip()
    return text

# ==============================
# Process Resume File
# ==============================

def parse_resume(file_path):
    """Parse the resume, extract structured text and emails, and format output."""
    if file_path.endswith('.pdf'):
        full_text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        full_text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are allowed.")

    # Extract emails
    emails = extract_emails(full_text)

    # Print formatted output
    print("\n" + "=" * 80)
    print(f"Processing File: {os.path.basename(file_path)}")
    print("=" * 80)
    print(f"Extracted Emails: {', '.join(emails)}")
    print("\nExtracted Full Text:\n" + "=" * 80)
    print(full_text if full_text else "No text extracted.")
    print("=" * 80 + "\n")

# ==============================
# Main Execution for Single File
# ==============================

if __name__ == "__main__":
    file_path = r"C:\Users\kkgau\Downloads\SrinivasaR.pdf"  # Change this to your file path
    if os.path.exists(file_path):
        parse_resume(file_path)
    else:
        print("File does not exist. Please check the path.")
