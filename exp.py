# import os
# import re
# import platform
# import mimetypes
# import tempfile
# import subprocess
# import hashlib

# import docx
# import pdfplumber

# def convert_doc_to_docx(doc_file):
#     """
#     Converts a .doc binary file to .docx format.
#     Supports:
#     - Microsoft Word (Windows)
#     - LibreOffice (Linux/macOS)
#     """
#     # Save binary file to a temporary .doc file
#     with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_doc:
#         temp_doc.write(doc_file.read())  # Save binary content to file
#         temp_doc_path = temp_doc.name  # Get the path
 
#     output_dir = os.path.dirname(temp_doc_path)
#     new_file = os.path.splitext(temp_doc_path)[0] + ".docx"
 
#     if platform.system() == "Windows":
#         try:
#             from comtypes.client import CreateObject
#             word = CreateObject("Word.Application")
#             word.Visible = False
#             doc = word.Documents.Open(temp_doc_path)
#             doc.SaveAs(new_file, 16)  # 16 represents DOCX format
#             doc.Close()
#             word.Quit()
#             print(f"✅ Converted {temp_doc_path} to {new_file}")
#         except Exception as e:
#             print(f"❌ Error converting {temp_doc_path}: {e}")
#             new_file = None
 
#     else:  # Linux/macOS (Docker)
#         soffice_cmd = "/usr/bin/soffice"
#         if not os.path.exists(soffice_cmd):
#             print("❌ LibreOffice (soffice) not found! Install it with sudo apt install libreoffice")
#             new_file = None
#         else:
#             try:
#                 subprocess.run(
#                     [soffice_cmd, '--headless', '--convert-to', 'docx', temp_doc_path, '--outdir', output_dir],
#                     stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#                 )
#                 if not os.path.exists(new_file):
#                     print(f"⚠️ Conversion executed but {new_file} was not found.")
#                     new_file = None
#             except subprocess.CalledProcessError as e:
#                 print(f"❌ Error converting {temp_doc_path}: {e.stderr}")
#                 new_file = None
 
#     os.remove(temp_doc_path)  # Cleanup temporary .doc file
#     return new_file
 
# def extract_text_from_pdf(pdf_path):
#     """Extracts text from a PDF while preserving formatting."""
#     text = ""
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text += page_text + "\n\n"
#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
#     return text.strip()
 
# def extract_text_from_docx(docx_path):
#     """Extracts text from a DOCX file, preserving paragraphs and tables."""
#     paragraphs = []
#     try:
#         doc = docx.Document(docx_path)
#         # Extract paragraphs
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 paragraphs.append(para.text.strip())
 
#         # Extract tables
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
#                 if row_text:
#                     paragraphs.append(row_text)
 
#     except Exception as e:
#         print(f"❌ Error reading DOCX file {docx_path}: {e}")
#     return "\n\n".join(paragraphs)
 
# def extract_resume(file_path):
#     extracted_text = ""
#     phone_pattern = re.compile(r'\b\d{10,15}\b')  # Detects valid phone numbers
#     email_pattern = re.compile(r'[\w\.-]+@[\w\.-]+')  # Detects valid emails
 
#     mime_type, _ = mimetypes.guess_type(file_path.name)
   
#     if mime_type == 'application/pdf':
#         extracted_text = extract_text_from_pdf(file_path)
   
#     elif mime_type in [
#         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     ]:
#         extracted_text = extract_text_from_docx(file_path)
   
#     elif mime_type in [
#         "application/msword"
#     ]:
#         # Create a temporary file for the .doc file
#         converted_docx_path = convert_doc_to_docx(file_path)  # Now works with binary file
       
#         if converted_docx_path:
#             extracted_text = extract_text_from_docx(converted_docx_path)
#             print("extracted_text", extracted_text, flush=True)
#             os.remove(converted_docx_path)  # Cleanup converted .docx file
#         else:
#             raise ValueError("❌ Failed to convert DOC file to DOCX.")
#         # file_path.seek(0)
#         # extracted_text = extract_text_from_doc(file_path)
               
#     extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()  # Normalize whitespace
#     # print("extracted_text", extracted_text, flush=True)
#     # Generate hash of the extracted text
#     resume_hash = hashlib.sha256(extracted_text.encode()).hexdigest()
#     print("resume_hash", resume_hash, flush=True)
#     return extracted_text, resume_hash
#################################################
import os
import re
import platform
import mimetypes
import tempfile
import subprocess
import hashlib
import docx
import pdfplumber

# def convert_doc_to_docx(doc_file_path):
#     # with open(doc_file_path, "rb") as doc_file:
#     #     with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_doc:
#     #         temp_doc.write(doc_file.read())
#     #         temp_doc_path = temp_doc.name

#     # output_dir = os.path.dirname(temp_doc_path)
#     # new_file = os.path.splitext(temp_doc_path)[0] + ".docx"
def convert_doc_to_docx(doc_file):
    """
    Converts a .doc binary file to .docx format.
    Supports:
    - Microsoft Word (Windows)
    - LibreOffice (Linux/macOS)
    """
    # Save binary file to a temporary .doc file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_doc:
        temp_doc.write(doc_file.read())  # Save binary content to file
        temp_doc_path = temp_doc.name  # Get the path
 
    output_dir = os.path.dirname(temp_doc_path)
    new_file = os.path.splitext(temp_doc_path)[0] + ".docx"
 

    if platform.system() == "Windows":
        try:
            from comtypes.client import CreateObject
            word = CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(temp_doc_path)
            doc.SaveAs(new_file, 16)
            doc.Close()
            word.Quit()
            print(f"✅ Converted {temp_doc_path} to {new_file}")
        except Exception as e:
            print(f"❌ Error converting {temp_doc_path}: {e}")
            new_file = None
    else:
        soffice_cmd = "/usr/bin/soffice"
        if not os.path.exists(soffice_cmd):
            print("❌ LibreOffice (soffice) not found! Install it with: sudo apt install libreoffice")
            new_file = None
        else:
            try:
                subprocess.run(
                    [soffice_cmd, '--headless', '--convert-to', 'docx', temp_doc_path, '--outdir', output_dir],
                    stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
                )
                if not os.path.exists(new_file):
                    print(f"⚠️ Conversion executed but {new_file} was not found.")
                    new_file = None
            except subprocess.CalledProcessError as e:
                print(f"❌ Error converting {temp_doc_path}: {e.stderr}")
                new_file = None

    os.remove(temp_doc_path)
    return new_file

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
    except Exception as e:
        print(f"❌ Error reading PDF file {pdf_path}: {e}")
    return text.strip()

def extract_text_from_docx(docx_path):
    paragraphs = []
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
    except Exception as e:
        print(f"❌ Error reading DOCX file {docx_path}: {e}")
    return "\n\n".join(paragraphs)

def extract_resume(file_path):
    extracted_text = ""
    phone_pattern = re.compile(r'\b\d{10,15}\b')  # Detects valid phone numbers
    email_pattern = re.compile(r'[\w\.-]+@[\w\.-]+')  # Detects valid emails
    mime_type, _ = mimetypes.guess_type(file_path)

    if mime_type == 'application/pdf':
        extracted_text = extract_text_from_pdf(file_path)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        extracted_text = extract_text_from_docx(file_path)
    elif mime_type == "application/msword":
        converted_docx_path = convert_doc_to_docx(file_path)
        if converted_docx_path:
            extracted_text = extract_text_from_docx(converted_docx_path)
            os.remove(converted_docx_path)
        else:
            raise ValueError("❌ Failed to convert DOC file to DOCX.")
    else:
        raise ValueError(f"❌ Unsupported file type: {mime_type}")

    extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
    return extracted_text
    # resume_hash = hashlib.sha256(extracted_text.encode()).hexdigest()
    # print("resume_hash", resume_hash, flush=True)
    # return extracted_text, resume_hash

# ------------------ ENTRY POINT ------------------

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("❌ Please provide a resume file path.")
        sys.exit(1)

    resume_path = sys.argv[1]

    if not os.path.exists(resume_path):
        print(f"❌ File not found: {resume_path}")
        sys.exit(1)

    resume_text = extract_resume(resume_path)
    print("\n📄 Extracted Resume Content:\n")
    print(resume_text)
