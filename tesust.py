# import os
# import subprocess
# import platform
# import pdfplumber
# import docx

# def convert_doc_to_docx(file_path):
#     """
#     Converts a .doc file to .docx format.
#     Uses:
#     - Microsoft Word (Windows)
#     - LibreOffice (Linux/macOS)
#     """
#     output_dir = os.path.dirname(file_path)
#     new_file = os.path.splitext(file_path)[0] + '.docx'

#     if platform.system() == "Windows":
#         try:
#             from comtypes.client import CreateObject  # Import only on Windows
#             word = CreateObject("Word.Application")
#             word.Visible = False
#             doc = word.Documents.Open(file_path)
#             doc.SaveAs(new_file, 16)  # 16 represents DOCX format
#             doc.Close()
#             word.Quit()
#             print(f"✅ Converted {file_path} to {new_file}")
#             return new_file
#         except Exception as e:
#             print(f"❌ Error converting {file_path}: {e}")
#             return None

#     else:  # Linux/macOS (Docker)
#         soffice_cmd = "/usr/bin/soffice"
#         if not os.path.exists(soffice_cmd):
#             print("❌ LibreOffice (`soffice`) not found! Install it with `sudo apt install libreoffice`")
#             return None
#         try:
#             result = subprocess.run(
#                 [soffice_cmd, '--headless', '--convert-to', 'docx', file_path, '--outdir', output_dir],
#                 stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#             )
#             if os.path.exists(new_file):
#                 print(f"✅ Converted {file_path} to {new_file}")
#                 return new_file
#             else:
#                 print(f"⚠️ Conversion executed but {new_file} was not found.")
#                 return None
#         except subprocess.CalledProcessError as e:
#             print(f"❌ Error converting {file_path}: {e.stderr}")
#             return None

# def extract_text_from_pdf(pdf_path):
#     """Extracts text from a PDF while preserving formatting."""
#     text = ""
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text += page_text + "\n\n"
#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
#     return text.strip()

# def extract_text_from_docx(docx_path):
#     """Extracts text from a DOCX file, preserving paragraphs and tables."""
#     paragraphs = []
#     try:
#         doc = docx.Document(docx_path)
#         # Extract paragraphs
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 paragraphs.append(para.text.strip())

#         # Extract tables
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
#                 if row_text:
#                     paragraphs.append(row_text)

#     except Exception as e:
#         print(f"❌ Error reading DOCX file {docx_path}: {e}")
#     return "\n\n".join(paragraphs)

# def extract_resume_text(file_path):
#     """Extracts text from PDF, DOCX, or converts DOC to DOCX before extracting."""
#     file_path_lower = file_path.lower()
#     if file_path_lower.endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path_lower.endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     elif file_path_lower.endswith(".doc"):
#         converted = convert_doc_to_docx(file_path)
#         if converted:
#             return extract_text_from_docx(converted)
#         else:
#             raise ValueError(f"❌ Failed to convert {file_path} to DOCX.")
#     else:
#         raise ValueError("❌ Unsupported file format. Only PDF, DOC, and DOCX are supported.")

# def process_resumes_in_folder(folder_path, limit=10):
#     """Processes up to `limit` resume files in the folder and prints extracted text."""
#     files = [f for f in os.listdir(folder_path) if f.lower().endswith(('.pdf', '.doc', '.docx'))]
#     files = files[:limit]

#     if not files:
#         print("❌ No resume files found in the folder.")
#         return

#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         print("=" * 80)
#         print(f"📂 Processing File: {filename}")
#         print("=" * 80)
#         try:
#             text = extract_resume_text(file_path)
#             if text:
#                 print(text)
#             else:
#                 print("⚠️ No text extracted from the file.")
#         except Exception as e:
#             print(f"❌ Error processing file {filename}: {e}")
#         print("=" * 80 + "\n")

# if __name__ == "__main__":
#     # Change this path based on your Docker mount
#     folder_path = "/app"
#     process_resumes_in_folder(folder_path, limit=10)
###################################################################
# import os
# import re
# import subprocess
# import pdfplumber
# import docx
# import textract
# from pdf2image import convert_from_path
# from pytesseract import image_to_string

# def convert_doc_to_docx(file_path):
#     """Converts a .doc file to .docx format using LibreOffice."""
#     output_dir = os.path.dirname(file_path)
#     new_file = os.path.splitext(file_path)[0] + '.docx'
#     soffice_cmd = "/usr/bin/soffice"
    
#     if not os.path.exists(soffice_cmd):
#         print("❌ LibreOffice (`soffice`) not found! Install it with `sudo apt install libreoffice`")
#         return None
#     try:
#         subprocess.run(
#             [soffice_cmd, '--headless', '--convert-to', 'docx', file_path, '--outdir', output_dir],
#             stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#         )
#         return new_file if os.path.exists(new_file) else None
#     except subprocess.CalledProcessError as e:
#         print(f"❌ Error converting {file_path}: {e.stderr}")
#         return None

# def extract_text_from_pdf(pdf_path):
#     """Extracts text from a PDF, with OCR fallback for scanned PDFs."""
#     text = ""
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 text += (page.extract_text() or "") + "\n"
#         if not text.strip():
#             text = perform_ocr(pdf_path)
#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
#     return text.strip()

# def extract_text_from_docx(file_path):
#     """Extracts text from a DOCX file, ensuring structured formatting."""
#     text = []
#     try:
#         doc = docx.Document(file_path)
        
#         # Extract headers
#         for section in doc.sections:
#             if section.header:
#                 text.extend(f"HEADER: {p.text.strip()}" for p in section.header.paragraphs if p.text.strip())
        
#         # Extract paragraphs, preserving bullet points
#         for para in doc.paragraphs:
#             stripped_text = para.text.strip()
#             if not stripped_text:
#                 continue
#             if para.style and para.style.name.startswith("Heading"):
#                 text.append(f"\n{stripped_text.upper()}\n")
#             elif stripped_text.startswith(("•", "-", "*")):
#                 text.append(f"  - {stripped_text}")
#             else:
#                 text.append(stripped_text)
        
#         # Extract tables
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 if row_text:
#                     text.append(row_text)
        
#         # Extract footers
#         for section in doc.sections:
#             if section.footer:
#                 text.extend(f"FOOTER: {p.text.strip()}" for p in section.footer.paragraphs if p.text.strip())
#     except Exception as e:
#         print(f"❌ Error reading DOCX file {file_path}: {e}")
#     return "\n".join(text)

# def perform_ocr(file_path):
#     """Performs OCR on scanned PDFs."""
#     try:
#         images = convert_from_path(file_path)
#         return "\n".join(image_to_string(img, lang="eng") for img in images).strip()
#     except Exception as e:
#         print(f"❌ OCR failed: {e}")
#         return ""

# def extract_emails(text):
#     """Extract email addresses from text."""
#     email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
#     return list(set(re.findall(email_pattern, text)))

# def extract_resume_text(file_path):
#     """Extracts text from PDF, DOCX, or converts DOC to DOCX before extracting."""
#     if file_path.lower().endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path.lower().endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     elif file_path.lower().endswith(".doc"):
#         converted = convert_doc_to_docx(file_path)
#         return extract_text_from_docx(converted) if converted else ""
#     else:
#         raise ValueError("❌ Unsupported file format. Only PDF, DOC, and DOCX are supported.")

# def process_resumes_in_folder(folder_path, limit=10):
#     """Processes up to `limit` resume files in the folder and prints extracted text."""
#     files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".doc", ".docx"))][:limit]
#     if not files:
#         print("❌ No resume files found in the folder.")
#         return
#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         print("=" * 80)
#         print(f"📂 Processing File: {filename}")
#         print("=" * 80)
#         try:
#             text = extract_resume_text(file_path)
#             emails = extract_emails(text)
#             print(f"Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}\n")
#             print(text if text else "⚠️ No text extracted from the file.")
#         except Exception as e:
#             print(f"❌ Error processing file {filename}: {e}")
#         print("=" * 80 + "\n")

# if __name__ == "__main__":
#     folder_path = "/app"
#     process_resumes_in_folder(folder_path, limit=10)
#####################################
# import os
# import re
# import subprocess
# import pdfplumber
# import docx
# import textwrap
# from pdf2image import convert_from_path
# from pytesseract import image_to_string

# # ============================
# # Convert DOC to DOCX
# # ============================
# def convert_doc_to_docx(file_path):
#     """Converts a .doc file to .docx format using LibreOffice."""
#     output_dir = os.path.dirname(file_path)
#     new_file = os.path.splitext(file_path)[0] + ".docx"
#     soffice_cmd = "/usr/bin/soffice"

#     if not os.path.exists(soffice_cmd):
#         print("❌ LibreOffice (`soffice`) not found! Install it with `sudo apt install libreoffice`")
#         return None
#     try:
#         subprocess.run(
#             [soffice_cmd, "--headless", "--convert-to", "docx", file_path, "--outdir", output_dir],
#             stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#         )
#         return new_file if os.path.exists(new_file) else None
#     except subprocess.CalledProcessError as e:
#         print(f"❌ Error converting {file_path}: {e.stderr}")
#         return None

# # ============================
# # Extract Text from PDF
# # ============================
# def extract_text_from_pdf(pdf_path):
#     """Extracts text from a PDF while preserving formatting."""
#     text = []
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text.append(page_text)
#         if not text:
#             text.append(perform_ocr(pdf_path))  # Perform OCR if no text is found
#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
#     return format_resume_text("\n\n".join(text))

# # ============================
# # Extract Text from DOCX with Formatting
# # ============================
# def extract_text_from_docx(file_path):
#     """Extracts text from a DOCX file while preserving layout and formatting."""
#     text = []
#     try:
#         doc = docx.Document(file_path)

#         # Extract headers
#         for section in doc.sections:
#             if section.header:
#                 text.extend(f"HEADER: {p.text.strip()}" for p in section.header.paragraphs if p.text.strip())

#         # Extract paragraphs and preserve formatting
#         for para in doc.paragraphs:
#             stripped_text = para.text.strip()
#             if not stripped_text:
#                 continue
#             if para.style and para.style.name.startswith("Heading"):
#                 text.append(f"\n{stripped_text.upper()}\n")  # Preserve headings
#             elif para.text.startswith(("•", "-", "*")):
#                 text.append(f"  - {stripped_text}")  # Maintain bullet points
#             else:
#                 text.append(stripped_text)

#         # Extract tables and format them as structured text
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 if row_text:
#                     text.append(row_text)

#         # Extract footers
#         for section in doc.sections:
#             if section.footer:
#                 text.extend(f"FOOTER: {p.text.strip()}" for p in section.footer.paragraphs if p.text.strip())
#     except Exception as e:
#         print(f"❌ Error reading DOCX file {file_path}: {e}")
#     return format_resume_text("\n".join(text))

# # ============================
# # Perform OCR on Scanned PDFs
# # ============================
# def perform_ocr(file_path):
#     """Performs OCR on scanned PDFs to extract text."""
#     try:
#         images = convert_from_path(file_path)
#         return "\n".join(image_to_string(img, lang="eng") for img in images).strip()
#     except Exception as e:
#         print(f"❌ OCR failed: {e}")
#         return ""

# # ============================
# # Format Resume Text Properly
# # ============================
# def format_resume_text(text, width=100):
#     """Formats extracted text to improve layout and readability."""
#     lines = text.split("\n")
#     formatted_lines = []

#     for line in lines:
#         stripped_line = line.strip()

#         if stripped_line.startswith(("•", "-", "*")):  # Preserve bullet points
#             formatted_lines.append(stripped_line)
#         elif stripped_line.isupper():  # Preserve uppercase section headers
#             formatted_lines.append(f"\n{stripped_line}\n")
#         else:
#             formatted_lines.append(textwrap.fill(stripped_line, width=width))

#     return "\n".join(formatted_lines)

# # ============================
# # Extract Emails
# # ============================
# def extract_emails(text):
#     """Extracts email addresses from the text."""
#     email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
#     return list(set(re.findall(email_pattern, text)))

# # ============================
# # Extract Resume Text
# # ============================
# def extract_resume_text(file_path):
#     """Extracts text from PDF, DOCX, or converts DOC to DOCX before extracting."""
#     if file_path.lower().endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path.lower().endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     elif file_path.lower().endswith(".doc"):
#         converted = convert_doc_to_docx(file_path)
#         return extract_text_from_docx(converted) if converted else ""
#     else:
#         raise ValueError("❌ Unsupported file format. Only PDF, DOC, and DOCX are supported.")

# # ============================
# # Process Resumes in Folder
# # ============================
# def process_resumes_in_folder(folder_path, limit=10):
#     """Processes up to `limit` resume files in the folder and prints extracted text."""
#     files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".doc", ".docx"))][:limit]
#     if not files:
#         print("❌ No resume files found in the folder.")
#         return
#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         print("=" * 100)
#         print(f"📂 Processing File: {filename}")
#         print("=" * 100)
#         try:
#             text = extract_resume_text(file_path)
#             emails = extract_emails(text)
#             print(f"📧 Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}\n")
#             print(text if text else "⚠️ No text extracted from the file.")
#         except Exception as e:
#             print(f"❌ Error processing file {filename}: {e}")
#         print("=" * 100 + "\n")

# # ============================
# # Main Execution
# # ============================
# if __name__ == "__main__":
#     folder_path = "/app"
#     process_resumes_in_folder(folder_path, limit=10)
############################################
# import os
# import re
# import subprocess
# import pdfplumber
# import docx
# import textwrap
# from pdf2image import convert_from_path
# from pytesseract import image_to_string

# # ============================
# # Convert DOC to DOCX
# # ============================
# def convert_doc_to_docx(file_path):
#     """Converts a .doc file to .docx format using LibreOffice."""
#     output_dir = os.path.dirname(file_path)
#     new_file = os.path.splitext(file_path)[0] + ".docx"
#     soffice_cmd = "/usr/bin/soffice"

#     if not os.path.exists(soffice_cmd):
#         print("❌ LibreOffice (soffice) not found! Install it with sudo apt install libreoffice")
#         return None
#     try:
#         subprocess.run(
#             [soffice_cmd, "--headless", "--convert-to", "docx", file_path, "--outdir", output_dir],
#             stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#         )
#         return new_file if os.path.exists(new_file) else None
#     except subprocess.CalledProcessError as e:
#         print(f"❌ Error converting {file_path}: {e.stderr}")
#         return None

# # ============================
# # Extract Text from PDF
# # ============================
# def extract_text_from_pdf(pdf_path):
#     """Extracts text from a PDF while preserving formatting."""
#     text = []
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text.append(page_text)
#         if not text:
#             text.append(perform_ocr(pdf_path))  # Perform OCR if no text is found
#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
#     return format_resume_text("\n\n".join(text))

# # ============================
# # Extract Text from DOCX with Formatting
# # ============================
# def extract_text_from_docx(file_path):
#     """Extracts text from a DOCX file while preserving layout and formatting."""
#     text = []
#     try:
#         doc = docx.Document(file_path)

#         # Extract headers
#         for section in doc.sections:
#             if section.header:
#                 text.extend(f"📌 {p.text.strip()}" for p in section.header.paragraphs if p.text.strip())

#         # Extract paragraphs and preserve formatting
#         for para in doc.paragraphs:
#             stripped_text = para.text.strip()
#             if not stripped_text:
#                 continue
#             if para.style and para.style.name.startswith("Heading"):
#                 text.append(f"\n📢 {stripped_text.upper()}\n")  # Preserve headings
#             elif para.text.startswith(("•", "-", "*")):
#                 text.append(f"  ➜ {stripped_text}")  # Maintain bullet points
#             else:
#                 text.append(stripped_text)

#         # Extract tables and format them as structured text
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 if row_text:
#                     text.append(row_text)

#         # Extract footers
#         for section in doc.sections:
#             if section.footer:
#                 text.extend(f"📌 {p.text.strip()}" for p in section.footer.paragraphs if p.text.strip())
#     except Exception as e:
#         print(f"❌ Error reading DOCX file {file_path}: {e}")
#     return format_resume_text("\n".join(text))

# # ============================
# # Perform OCR on Scanned PDFs
# # ============================
# def perform_ocr(file_path):
#     """Performs OCR on scanned PDFs and extracts text for hidden emails."""
#     try:
#         images = convert_from_path(file_path)
#         text = "\n".join(image_to_string(img, lang="eng") for img in images).strip()
#         return text
#     except Exception as e:
#         print(f"❌ OCR failed: {e}")
#         return ""

# # ============================
# # Format Resume Text Properly
# # ============================
# def format_resume_text(text, width=100):
#     """Formats extracted text to improve layout and readability."""
#     lines = text.split("\n")
#     formatted_lines = []
#     is_heading = False

#     for line in lines:
#         stripped_line = line.strip()

#         # Handle Bullet Points
#         if stripped_line.startswith(("•", "-", "*")):
#             formatted_lines.append(f"  ➜ {stripped_line}")

#         # Preserve Section Headers
#         elif stripped_line.isupper():
#             formatted_lines.append(f"\n📢 {stripped_line}\n")
#             is_heading = True

#         # Ensure Proper Paragraph Formatting
#         elif is_heading:
#             formatted_lines.append(stripped_line)
#             is_heading = False
#         else:
#             formatted_lines.append(textwrap.fill(stripped_line, width=width))

#     return "\n".join(formatted_lines)

# # ============================
# # Extract Emails
# # ============================
# def extract_emails(text):
#     """Extracts email addresses from the text, including 'mailto:' links."""
#     email_pattern = r"(?:mailto:)?([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})"
#     return list(set(re.findall(email_pattern, text)))

# # ============================
# # Extract Resume Text
# # ============================
# def extract_resume_text(file_path):
#     """Extracts text from PDF, DOCX, or converts DOC to DOCX before extracting."""
#     if file_path.lower().endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path.lower().endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     elif file_path.lower().endswith(".doc"):
#         converted = convert_doc_to_docx(file_path)
#         return extract_text_from_docx(converted) if converted else ""
#     else:
#         raise ValueError("❌ Unsupported file format. Only PDF, DOC, and DOCX are supported.")

# # ============================
# # Process Resumes in Folder
# # ============================
# def process_resumes_in_folder(folder_path, limit=10):
#     """Processes up to limit resume files in the folder and extracts emails."""
#     files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".doc", ".docx"))][:limit]
#     if not files:
#         print("❌ No resume files found in the folder.")
#         return
    
#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         print("=" * 100)
#         print(f"📂 Processing File: {filename}")
#         print("=" * 100)
#         try:
#             text = extract_resume_text(file_path)
#             emails = extract_emails(text)
#             print(f"📧 Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}\n")
#             print(text if text else "⚠️ No text extracted from the file.")
#         except Exception as e:
#             print(f"❌ Error processing file {filename}: {e}")
#         print("=" * 100 + "\n")

# # ============================
# # Main Execution
# # ============================
# if __name__ == "__main__":
#     folder_path = "/app"
#     process_resumes_in_folder(folder_path, limit=10)
###########################################################
# import os
# import re
# import subprocess
# import pdfplumber
# import docx
# import textwrap
# from pdf2image import convert_from_path
# from pytesseract import image_to_string

# def convert_doc_to_docx(file_path):
#     output_dir = os.path.dirname(file_path)
#     new_file = os.path.splitext(file_path)[0] + ".docx"
#     soffice_cmd = "/usr/bin/soffice"
#     if not os.path.exists(soffice_cmd):
#         print("❌ LibreOffice (`soffice`) not found! Install it with `sudo apt install libreoffice`")
#         return None
#     try:
#         subprocess.run(
#             [soffice_cmd, "--headless", "--convert-to", "docx", file_path, "--outdir", output_dir],
#             stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#         )
#         return new_file if os.path.exists(new_file) else None
#     except subprocess.CalledProcessError as e:
#         print(f"❌ Error converting {file_path}: {e.stderr}")
#         return None

# def extract_text_from_pdf(pdf_path):
#     text = []
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text.append(page_text)
#         if not text:
#             print(f"🔍 No direct text found in {pdf_path}, performing OCR...")
#             text.append(perform_ocr(pdf_path))
#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
#     return format_resume_text("\n".join(text))

# def extract_text_from_docx(file_path):
#     text = []
#     try:
#         doc = docx.Document(file_path)
#         for section in doc.sections:
#             if section.header:
#                 text.extend(f"📌 {p.text.strip()}" for p in section.header.paragraphs if p.text.strip())
#         for para in doc.paragraphs:
#             stripped_text = para.text.strip()
#             if not stripped_text:
#                 continue
#             if para.style and para.style.name.startswith("Heading"):
#                 text.append(f"\n📢 {stripped_text.upper()}\n")
#             elif para.text.startswith(("•", "-", "*")):
#                 text.append(f"  ➜ {stripped_text}")
#             else:
#                 text.append(stripped_text)
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 if row_text:
#                     text.append(row_text)
#         for section in doc.sections:
#             if section.footer:
#                 text.extend(f"📌 {p.text.strip()}" for p in section.footer.paragraphs if p.text.strip())
#     except Exception as e:
#         print(f"❌ Error reading DOCX file {file_path}: {e}")
#     return format_resume_text("\n".join(text))

# def perform_ocr(file_path):
#     try:
#         images = convert_from_path(file_path)
#         text = "\n".join(image_to_string(img, lang="eng") for img in images).strip()
#         return text
#     except Exception as e:
#         print(f"❌ OCR failed: {e}")
#         return ""

# def format_resume_text(text, width=100):
#     lines = text.split("\n")
#     formatted_lines = []
#     is_heading = False
#     for line in lines:
#         stripped_line = line.strip()
#         if stripped_line.startswith(("•", "-", "*")):
#             formatted_lines.append(f"  ➜ {stripped_line}")
#         elif stripped_line.isupper():
#             formatted_lines.append(f"\n📢 {stripped_line}\n")
#             is_heading = True
#         elif is_heading:
#             formatted_lines.append(stripped_line)
#             is_heading = False
#         else:
#             formatted_lines.append(textwrap.fill(stripped_line, width=width))
#     return "\n".join(formatted_lines)

# def extract_emails(text):
#     text = text.replace("\n", " ").replace("\r", " ")
#     email_pattern = r"(?:mailto:)?([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-zA-Z]{2,})"
#     emails = re.findall(email_pattern, text)
#     return list(set(emails))

# def extract_urls(text):
#     url_pattern = r"(https?://[^\s]+)"
#     urls = re.findall(url_pattern, text)
#     return list(set(urls))

# def extract_resume_text(file_path):
#     if file_path.lower().endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path.lower().endswith(".docx"):
#         return extract_text_from_docx(file_path)
#     elif file_path.lower().endswith(".doc"):
#         converted = convert_doc_to_docx(file_path)
#         return extract_text_from_docx(converted) if converted else ""
#     else:
#         raise ValueError("❌ Unsupported file format. Only PDF, DOC, and DOCX are supported.")

# def process_resumes_in_folder(folder_path, limit=10):
#     files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".doc", ".docx"))][:limit]
#     if not files:
#         print("❌ No resume files found in the folder.")
#         return
#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         print("=" * 100)
#         print(f"📂 Processing File: {filename}")
#         print("=" * 100)
#         try:
#             text = extract_resume_text(file_path)
#             emails = extract_emails(text)
#             urls = extract_urls(text)
#             print(f"📧 Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}")
#             print(f"🔗 Extracted URLs: {', '.join(urls) if urls else 'No URLs found.'}\n")
#             print(text[:2000])
#         except Exception as e:
#             print(f"❌ Error processing file {filename}: {e}")
#         print("=" * 100 + "\n")

# if __name__ == "__main__":
#     folder_path = "/app"
#     process_resumes_in_folder(folder_path, limit=10)
######################
import os
import re
import subprocess
import pdfplumber
import docx
import textwrap
from pdf2image import convert_from_path
from pytesseract import image_to_string

def convert_doc_to_docx(file_path):
    output_dir = os.path.dirname(file_path)
    new_file = os.path.splitext(file_path)[0] + ".docx"
    soffice_cmd = "/usr/bin/soffice"

    if not os.path.exists(soffice_cmd):
        print("LibreOffice (`soffice`) not found! Install it with `sudo apt install libreoffice`")
        return None
    try:
        subprocess.run(
            [soffice_cmd, "--headless", "--convert-to", "docx", file_path, "--outdir", output_dir],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
        )
        return new_file if os.path.exists(new_file) else None
    except subprocess.CalledProcessError as e:
        print(f"Error converting {file_path}: {e.stderr}")
        return None

def extract_text_from_binary(binary_data, file_extension):
    temp_filename = "temp_file." + file_extension
    with open(temp_filename, "wb") as temp_file:
        temp_file.write(binary_data)
    
    if file_extension == "pdf":
        text, emails = extract_text_from_pdf(temp_filename)
    elif file_extension == "docx":
        text, emails = extract_text_from_docx(temp_filename)
    elif file_extension == "doc":
        converted = convert_doc_to_docx(temp_filename)
        text, emails = extract_text_from_docx(converted) if converted else ("", [])
    else:
        raise ValueError("Unsupported file format. Only PDF, DOC, and DOCX are supported.")
    
    os.remove(temp_filename)
    return text, emails

def extract_text_from_pdf(pdf_path):
    text = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        if not text:
            print(f"No direct text found in {pdf_path}, performing OCR...")
            text.append(perform_ocr(pdf_path))
    except Exception as e:
        print(f"Error reading PDF file {pdf_path}: {e}")
    raw_text = "\n".join(text)
    emails = extract_emails(raw_text)
    return format_resume_text(raw_text), emails

def extract_text_from_docx(file_path):
    text = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            stripped_text = para.text.strip()
            if stripped_text:
                text.append(stripped_text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text.append(row_text)
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
    raw_text = "\n".join(text)
    emails = extract_emails(raw_text)
    return format_resume_text(raw_text), emails

def perform_ocr(file_path):
    try:
        images = convert_from_path(file_path)
        text = "\n".join(image_to_string(img, lang="eng") for img in images).strip()
        return text
    except Exception as e:
        print(f"OCR failed: {e}")
        return ""

def format_resume_text(text, width=100):
    lines = text.split("\n")
    formatted_lines = []
    for line in lines:
        stripped_line = line.strip()
        formatted_lines.append(textwrap.fill(stripped_line, width=width))
    return "\n".join(formatted_lines)

def extract_emails(text):
    text = text.replace("\n", " ").replace("\r", " ")
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    emails = re.findall(email_pattern, text)
    return list(set(emails))

def convert_files_to_binary(folder_path):
    file_binaries = {}
    files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".doc", ".docx"))]
    for filename in files:
        file_path = os.path.join(folder_path, filename)
        with open(file_path, "rb") as f:
            file_binaries[filename] = f.read()
    return file_binaries

def process_all_resumes(folder_path):
    file_binaries = convert_files_to_binary(folder_path)
    for filename, binary_data in file_binaries.items():
        file_extension = filename.split(".")[-1].lower()
        print("=" * 100)
        print(f"Processing File: {filename}")
        print("=" * 100)
        try:
            text, emails = extract_text_from_binary(binary_data, file_extension)
            print(f"Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}")
            print(text)
        except Exception as e:
            print(f"Error processing {filename}: {e}")
        print("=" * 100 + "\n")

if __name__ == "__main__":
    folder_path = os.getcwd()
    process_all_resumes(folder_path)

###############################################################################
# import os
# import re
# import subprocess
# import pdfplumber
# import docx
# import base64
# from pdf2image import convert_from_path
# from pytesseract import image_to_string

# def convert_doc_to_docx(file_path):
#     output_dir = os.path.dirname(file_path)
#     new_file = os.path.splitext(file_path)[0] + ".docx"
#     soffice_cmd = "/usr/bin/soffice"
    
#     if not os.path.exists(soffice_cmd):
#         print("❌ LibreOffice (`soffice`) not found! Install it with `sudo apt install libreoffice`")
#         return None
#     try:
#         subprocess.run(
#             [soffice_cmd, "--headless", "--convert-to", "docx", file_path, "--outdir", output_dir],
#             stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True
#         )
#         return new_file if os.path.exists(new_file) else None
#     except subprocess.CalledProcessError as e:
#         print(f"❌ Error converting {file_path}: {e.stderr}")
#         return None

# def extract_text_from_pdf(pdf_path):
#     text = ""
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 text += (page.extract_text() or "") + "\n"
#         if not text.strip():
#             text = perform_ocr(pdf_path)
#     except Exception as e:
#         print(f"❌ Error reading PDF file {pdf_path}: {e}")
#     return text.strip()

# def extract_text_from_docx(file_path):
#     text = []
#     try:
#         doc = docx.Document(file_path)
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 text.append(para.text.strip())
#         for table in doc.tables:
#             for row in table.rows:
#                 row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
#                 if row_text:
#                     text.append(row_text)
#     except Exception as e:
#         print(f"❌ Error reading DOCX file {file_path}: {e}")
#     return "\n".join(text)

# def perform_ocr(file_path):
#     try:
#         images = convert_from_path(file_path)
#         return "\n".join(image_to_string(img, lang="eng") for img in images).strip()
#     except Exception as e:
#         print(f"❌ OCR failed: {e}")
#         return ""

# def extract_emails(text):
#     email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
#     return list(set(re.findall(email_pattern, text)))

# def convert_file_to_binary(file_path):
#     try:
#         with open(file_path, "rb") as file:
#             return base64.b64encode(file.read()).decode("utf-8")
#     except Exception as e:
#         print(f"❌ Error converting {file_path} to binary: {e}")
#         return None

# def extract_resume_text(file_path):
#     binary_data = convert_file_to_binary(file_path)
#     if not binary_data:
#         return "", []

#     if file_path.lower().endswith(".pdf"):
#         text = extract_text_from_pdf(file_path)
#     elif file_path.lower().endswith(".docx"):
#         text = extract_text_from_docx(file_path)
#     elif file_path.lower().endswith(".doc"):
#         converted = convert_doc_to_docx(file_path)
#         text = extract_text_from_docx(converted) if converted else ""
#     else:
#         raise ValueError("❌ Unsupported file format. Only PDF, DOC, and DOCX are supported.")

#     emails = extract_emails(text)
#     return text, emails

# def process_resumes_in_folder(folder_path):
#     files = [f for f in os.listdir(folder_path) if f.lower().endswith((".pdf", ".doc", ".docx"))]
#     if not files:
#         print("❌ No resume files found in the folder.")
#         return

#     for filename in files:
#         file_path = os.path.join(folder_path, filename)
#         print("=" * 100)
#         print(f"📂 Processing File: {filename}")
#         print("=" * 100)

#         try:
#             text, emails = extract_resume_text(file_path)
#             print(f"📧 Extracted Emails: {', '.join(emails) if emails else 'No emails found.'}")
#             print("📜 Extracted Text:\n" + "=" * 100)
#             print(text)
#             print("\n" + "=" * 100)
#         except Exception as e:
#             print(f"❌ Error processing file {filename}: {e}")

# if __name__ == "__main__":
#     folder_path = os.getcwd()
#     process_resumes_in_folder(folder_path)
